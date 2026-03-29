"""
Menu Import Parser — extracts products from menu files using OpenAI Vision.

Supported formats:
  Images : .jpg .jpeg .png .gif .webp  → OpenAI Vision (gpt-4o)
  PDF    : .pdf (multi-page)           → PyMuPDF renders pages → Vision
  Word   : .docx                       → python-docx text → GPT-4o
  Excel  : .xlsx .xls                  → openpyxl / xlrd → direct mapping → GPT-4o fallback
  CSV    : .csv                         → direct mapping → GPT-4o fallback
  Text   : .txt                         → direct text → GPT-4o

Confidence scoring:
  1.0  = perfectly clear entry
  0.7+ = high confidence, ready to import
  0.5–0.7 = medium — flagged for review
  <0.5 = low — flagged, user must approve individually
"""

import os
import re
import json
import uuid
import base64
import logging
import tempfile
from pathlib import Path
from typing import List, Dict, Any, Optional

logger = logging.getLogger("menu-parser")

# Items with confidence below this threshold are flagged needs_review=True
REVIEW_THRESHOLD = 0.70

SUPPORTED_IMAGES = {".jpg", ".jpeg", ".png", ".gif", ".webp"}
SUPPORTED_DOCS   = {".pdf", ".docx", ".txt", ".csv", ".xlsx", ".xls"}
ALL_SUPPORTED    = SUPPORTED_IMAGES | SUPPORTED_DOCS


# ── OpenAI client ──────────────────────────────────────────────────────────────

def _get_client():
    key = os.getenv("OPENAI_API_KEY", "")
    if not key:
        raise ValueError("OPENAI_API_KEY غير مضبوط. أضفه في ملف .env")
    import openai
    return openai.OpenAI(api_key=key)


# ── Image → base64 ────────────────────────────────────────────────────────────

def _img_b64(path: str) -> str:
    ext = Path(path).suffix.lower().lstrip(".")
    mime_map = {"jpg": "jpeg", "jpeg": "jpeg", "png": "png", "gif": "gif", "webp": "webp"}
    mime = f"image/{mime_map.get(ext, 'jpeg')}"
    with open(path, "rb") as f:
        return f"data:{mime};base64,{base64.b64encode(f.read()).decode()}"


# ── Prompts ────────────────────────────────────────────────────────────────────

_VISION_PROMPT = """You are a professional restaurant menu parser. Carefully analyze this menu image.

Return ONLY a valid JSON array — no markdown, no explanation outside the array.

Each object in the array must have EXACTLY these fields:
{
  "name": "product name — keep original language (Arabic or English or both)",
  "category": "section / category name — keep original language",
  "price": number or null,
  "description": "short description or empty string",
  "variants": [{"name": "size or option", "price": null}],
  "confidence": 0.0–1.0,
  "needs_review": true or false,
  "source_note": "any parsing issue, or empty string"
}

STRICT RULES — violating these is unacceptable:
1. NEVER invent, guess, or estimate a price. If price text is unclear or missing → price: null.
2. NEVER invent product names. Only extract what is clearly visible.
3. confidence = 1.0 means 100% certain. Lower it for any ambiguity.
4. Set needs_review = true if: price is null, text is partially obscured, or meaning is ambiguous.
5. Preserve original Arabic and/or English text exactly as written.
6. If sizes/options exist (صغير/وسط/كبير or S/M/L), add them as variants.
7. price field is a number only — NO currency symbols.
8. Skip decorative headers or non-product lines.
9. If you see the same item mentioned multiple times, include it once.
10. Return the JSON array ONLY."""

_TEXT_PROMPT = """You are a professional restaurant menu parser. Extract all menu items from this text.

Return ONLY a valid JSON array — no markdown, no explanation outside the array.

Each object must have EXACTLY these fields:
{
  "name": "product name",
  "category": "section or category",
  "price": number or null,
  "description": "description or empty string",
  "variants": [{"name": "option name", "price": null}],
  "confidence": 0.0–1.0,
  "needs_review": true or false,
  "source_note": ""
}

Rules:
- DO NOT invent prices or names
- confidence = 1.0 for clearly structured data, lower for ambiguous entries
- If price is missing or unclear → price: null, needs_review: true
- Preserve original Arabic/English text
- Return ONLY the JSON array"""


# ── OpenAI calls ───────────────────────────────────────────────────────────────

def _vision_model() -> str:
    """Vision requires gpt-4o or gpt-4o-mini — force vision-capable model."""
    m = os.getenv("OPENAI_MODEL", "gpt-4o-mini")
    # Allow gpt-4o family; anything else falls back to gpt-4o-mini
    if m.startswith("gpt-4o"):
        return m
    return "gpt-4o-mini"


def _call_vision(image_paths: List[str], client) -> List[Dict]:
    """Send up to 4 images to gpt-4o Vision and extract menu items."""
    content = []
    for path in image_paths[:4]:
        try:
            content.append({
                "type": "image_url",
                "image_url": {"url": _img_b64(path), "detail": "high"}
            })
        except Exception as e:
            logger.warning(f"Image encode failed {path}: {e}")

    if not content:
        logger.warning("No images to send to Vision API")
        return []

    content.append({"type": "text", "text": _VISION_PROMPT})
    model = _vision_model()
    logger.info(f"Calling Vision API with model={model}, images={len(content)-1}")

    try:
        resp = client.chat.completions.create(
            model=model,
            messages=[{"role": "user", "content": content}],
            max_tokens=4000,
            temperature=0,
        )
        raw_text = resp.choices[0].message.content
        logger.info(f"Vision API response length={len(raw_text)}, preview={raw_text[:100]}")
        return _parse_json_response(raw_text)
    except Exception as e:
        err_str = str(e)
        logger.error(f"Vision API error (model={model}): {err_str}")
        if "insufficient_quota" in err_str or "429" in err_str:
            raise RuntimeError("OpenAI quota exceeded — أضف billing على https://platform.openai.com/account/billing")
        raise


def _call_text(text: str, client) -> List[Dict]:
    """Send text (up to 12k chars) to gpt-4o and extract menu items."""
    text = text[:12000].strip()
    if not text:
        return []
    model = os.getenv("OPENAI_MODEL", "gpt-4o-mini")
    logger.info(f"Calling Text API with model={model}, text_len={len(text)}")
    try:
        resp = client.chat.completions.create(
            model=model,
            messages=[{"role": "user", "content": f"{_TEXT_PROMPT}\n\nMenu text:\n{text}"}],
            max_tokens=4000,
            temperature=0,
        )
        raw_text = resp.choices[0].message.content
        logger.info(f"Text API response preview={raw_text[:100]}")
        return _parse_json_response(raw_text)
    except Exception as e:
        err_str = str(e)
        logger.error(f"Text API error: {err_str}")
        if "insufficient_quota" in err_str or "429" in err_str:
            raise RuntimeError("OpenAI quota exceeded — أضف billing على https://platform.openai.com/account/billing")
        raise


def _parse_json_response(raw: str) -> List[Dict]:
    """Strip markdown fences and parse JSON array from OpenAI response."""
    raw = raw.strip()
    # Strip ```json ... ``` fences
    if raw.startswith("```"):
        parts = raw.split("```")
        for part in parts:
            part = part.strip()
            if part.startswith("json"):
                part = part[4:].strip()
            if part.startswith("[") or part.startswith("{"):
                raw = part
                break
    try:
        result = json.loads(raw)
        return result if isinstance(result, list) else []
    except json.JSONDecodeError as e:
        logger.error(f"JSON parse error: {e} | raw[:200]={raw[:200]}")
        return []


# ── File type parsers ──────────────────────────────────────────────────────────

def _parse_pdf(path: str, client) -> List[Dict]:
    """Render each PDF page to a PNG image and extract via Vision."""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        logger.warning("PyMuPDF not installed — trying text extraction fallback")
        # Fallback: extract embedded text from PDF
        try:
            import fitz
        except Exception:
            return _error_item("PyMuPDF غير مثبت. نفّذ: pip install PyMuPDF")

    doc = fitz.open(path)
    all_items: List[Dict] = []
    batch: List[str] = []
    tmp_files: List[str] = []

    try:
        for page_num in range(len(doc)):
            page = doc[page_num]
            mat  = fitz.Matrix(2.0, 2.0)   # 2× zoom for better OCR quality
            pix  = page.get_pixmap(matrix=mat)
            tmp  = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
            pix.save(tmp.name)
            tmp.close()
            tmp_files.append(tmp.name)
            batch.append(tmp.name)

            # Process in batches of 2 pages to stay within token limits
            if len(batch) >= 2 or page_num == len(doc) - 1:
                items = _call_vision(batch, client)
                for item in items:
                    item.setdefault("source_page", page_num + 1)
                all_items.extend(items)
                for f in batch:
                    try: os.unlink(f)
                    except Exception: pass
                batch = []
    finally:
        doc.close()
        for f in tmp_files:
            try: os.unlink(f)
            except Exception: pass

    return all_items


def _parse_docx(path: str, client) -> List[Dict]:
    """Extract text + tables from .docx and parse with GPT-4o."""
    try:
        from docx import Document
        doc  = Document(path)
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        for table in doc.tables:
            for row in table.rows:
                text += "\n" + " | ".join(cell.text.strip() for cell in row.cells)
        return _call_text(text, client)
    except ImportError:
        return _error_item("python-docx غير مثبت. نفّذ: pip install python-docx")
    except Exception as e:
        logger.error(f"DOCX error: {e}")
        return []


# ╔══════════════════════════════════════════════════════════════════════════════╗
# ║               SPREADSHEET PARSER — DEFINITIONS & CONSTANTS                  ║
# ╚══════════════════════════════════════════════════════════════════════════════╝
#
# Column synonym sets are stored in normalised form (lowercase, separators
# removed) to be matched after _norm_header() is applied to the actual cell.

# ── Name column synonyms ──────────────────────────────────────────────────────
_COL_NAME = {
    # English
    "name", "productname", "itemname", "product", "item", "title",
    "dish", "meal", "menuitem", "fooditem", "food", "dishname", "foodname",
    "itemtitle", "mealdescription",
    # Arabic
    "اسم", "الاسم", "المنتج", "اسمالمنتج", "اسمالصنف", "الصنف",
    "الوجبة", "اسمالوجبة", "الطبق", "اسمالطبق", "العنصر",
    "اسمالعنصر", "عنصرالقائمة",
}

# ── Category column synonyms ──────────────────────────────────────────────────
_COL_CATEGORY = {
    # English
    "category", "cat", "section", "group", "type", "class",
    "subsection", "department", "course",
    # Arabic
    "قسم", "فئة", "القسم", "الفئة", "التصنيف", "النوع", "المجموعة",
    "التقسيم", "الفئةالرئيسية", "قسمالمنتج",
}

# ── Price column synonyms ─────────────────────────────────────────────────────
_COL_PRICE = {
    # English
    "price", "cost", "amount", "rate", "value", "unitprice",
    "listprice", "sellingprice", "saleprice", "retail", "retailprice",
    # Arabic
    "سعر", "السعر", "الثمن", "ثمن", "قيمة", "التكلفة", "سعرالوحدة",
    "سعربيع", "سعرالبيع", "سعرالتجزئة", "المبلغ", "التكلفه",
}

# ── Description column synonyms ───────────────────────────────────────────────
_COL_DESC = {
    # English
    "description", "desc", "details", "notes", "note", "info",
    "ingredients", "about", "summary",
    # Arabic
    "وصف", "الوصف", "تفاصيل", "ملاحظات", "ملاحظة",
    "المكونات", "مكونات", "معلومات",
}

# ── Icon / image column synonyms ──────────────────────────────────────────────
_COL_ICON  = {"icon", "emoji", "image", "img", "photo", "أيقونة", "رمز", "صورة"}

# ── Availability column synonyms ──────────────────────────────────────────────
_COL_AVAIL = {
    "available", "active", "enabled", "status", "instock", "visible",
    "متاح", "نشط", "الحالة", "مفعل", "متوفر",
}

# ── Row names that indicate non-product rows (normalised) ─────────────────────
_SKIP_ROW_NAMES = {
    # English
    "total", "subtotal", "grandtotal", "discount", "tax", "vat",
    "note", "notes", "remark", "remarks", "sum", "footer", "header",
    "page", "pagetotal", "nettotal",
    # Arabic
    "المجموع", "الإجمالي", "مجموع", "إجمالي", "ملاحظة", "ملاحظات",
    "خصم", "ضريبة", "ضريبةالقيمةالمضافة", "صفحة",
}

# ── Tokens that block partial header matching (ID/code columns, not name) ─────
_PARTIAL_MATCH_BLOCK = {
    "code", "id", "no", "num", "number", "ref", "sku",
    "barcode", "serial", "seq", "sequence", "index", "rank",
    "كود", "رقم", "مسلسل",
}

# ── Currency symbols to strip before numeric parsing (longest first) ──────────
_CURRENCY_SYMS = [
    "IQD", "EGP", "AED", "BHD", "QAR", "JOD", "LBP", "KWD", "OMR", "YER",
    "SAR", "MAD", "TND", "DZD", "SDG", "USD", "EUR", "GBP",
    "ر.س", "ريال", "ر.ع", "د.ك", "د.ع", "ج.م", "د.إ", "د.أ", "ل.ل", "ر.ق",
    "$", "€", "£", "﷼", "¥", "₺",
]

# Direct-parse quality threshold — below this score we also try OpenAI
_SHEET_QUALITY_FALLBACK = 0.30


# ╔══════════════════════════════════════════════════════════════════════════════╗
# ║                        NORMALISATION HELPERS                                ║
# ╚══════════════════════════════════════════════════════════════════════════════╝

def _norm_header(h: str) -> str:
    """Normalise a cell for column matching.
    Strips BOM / zero-width / control chars → lowercase → removes separators."""
    import unicodedata
    # Remove BOM and invisible Unicode characters
    h = h.replace('\ufeff', '').replace('\u200b', '').replace('\u200c', '')
    h = h.replace('\u200d', '').replace('\xa0', ' ')
    # Drop control (Cc) and format (Cf) characters, keep all printable Unicode
    h = ''.join(c for c in h if unicodedata.category(c) not in ('Cc', 'Cf'))
    h = h.strip().lower()
    # Remove separator characters so "product_name" == "productname" == "product name"
    for sep in (' ', '_', '-', '.', '/', '\\', '(', ')', '[', ']', ':'):
        h = h.replace(sep, '')
    return h


def _clean_cell(v: str) -> str:
    """Strip BOM, zero-width spaces, and control chars from any cell value.
    Preserves Arabic / Unicode text content completely."""
    import unicodedata
    v = v.replace('\ufeff', '').replace('\u200b', '').replace('\u200c', '')
    v = v.replace('\u200d', '').replace('\u200e', '').replace('\u200f', '')
    v = v.replace('\xa0', ' ')
    # Remove only actual control characters (Cc); keep all printable Unicode
    v = ''.join(c for c in v if unicodedata.category(c) != 'Cc')
    return v.strip()


def _header_to_field(hn: str) -> Optional[str]:
    """Map a normalised header string to a field key, or None if unrecognised.

    Strategy:
    1. Exact match against all synonym sets.
    2. Blocklist: if header contains an ID/code token, skip partial matching.
    3. Controlled substring match (keyword ≥ 4 chars to avoid noise).
    """
    pairs = (
        (_COL_NAME,     "name"),
        (_COL_CATEGORY, "category"),
        (_COL_PRICE,    "price"),
        (_COL_DESC,     "description"),
        (_COL_ICON,     "icon"),
        (_COL_AVAIL,    "available"),
    )
    # ── Exact match ───────────────────────────────────────────────────────────
    for col_set, field in pairs:
        if hn in col_set:
            return field

    # ── Blocklist: refuse partial match for ID/code/serial columns ────────────
    for block_tok in _PARTIAL_MATCH_BLOCK:
        if block_tok in hn:
            return None

    # ── Controlled substring / partial match ──────────────────────────────────
    # Keyword must be ≥ 4 chars; header must be ≥ 3 chars to reduce false positives
    for col_set, field in pairs[:4]:   # only name / category / price / description
        for kw in col_set:
            if len(kw) >= 4 and len(hn) >= 3 and (kw in hn or hn in kw):
                return field
    return None


# ── Comprehensive price parser ────────────────────────────────────────────────

def _parse_price(raw: str) -> Optional[float]:
    """Parse a price string to float.

    Handles:
    - Eastern Arabic numerals (٠١٢٣٤٥٦٧٨٩) and separators (٫ decimal, ٬ thousands)
    - All common currency symbols and codes (SAR, IQD, EGP, AED, د.ع, ر.س, etc.)
    - Smart comma: "5,000" → 5000.0 (thousands) vs "12,5" → 12.5 (decimal)
    - Broken formats like "5..0", multiple spaces, stray characters

    Returns None on failure or for negative values.
    """
    if not raw or not raw.strip():
        return None

    # Translate Eastern Arabic numerals and separator chars to ASCII equivalents
    # ٫ = Arabic decimal separator, ٬ = Arabic thousands separator
    _ea = str.maketrans("٠١٢٣٤٥٦٧٨٩٫٬", "0123456789.,")
    s = raw.strip().translate(_ea)

    # Strip currency symbols (longest first to avoid partial strip)
    for sym in _CURRENCY_SYMS:
        s = s.replace(sym, "")

    s = s.strip()
    if not s:
        return None

    # Smart comma disambiguation:
    #   "5,000"     → thousands (comma before exactly 3 digits at end) → 5000
    #   "1,234,567" → thousands → 1234567
    #   "12,5"      → decimal  → 12.5
    #   "12,50"     → decimal  → 12.50
    thousands_re = re.compile(r'^\d{1,3}(,\d{3})+(\.\d+)?$')
    decimal_comma_re = re.compile(r'^\d+,\d{1,2}$')

    if thousands_re.match(s):
        s = s.replace(',', '')          # "5,000" → "5000"
    elif decimal_comma_re.match(s):
        s = s.replace(',', '.')         # "12,5"  → "12.5"
    else:
        # Default: remaining commas treated as decimal separator
        s = s.replace(',', '.')

    # Strip non-numeric characters except dot and minus
    s = re.sub(r'[^\d.\-]', '', s)

    # Collapse multiple consecutive dots (e.g. "5..0" → "5.0")
    while '..' in s:
        s = s.replace('..', '.')

    # Handle edge cases like lone "." or "-"
    if not s or s in ('.', '-', '-.', '0.'):
        return None

    try:
        val = float(s)
        return val if val >= 0 else None
    except ValueError:
        return None


# ╔══════════════════════════════════════════════════════════════════════════════╗
# ║                        LOW-LEVEL FILE READERS                               ║
# ╚══════════════════════════════════════════════════════════════════════════════╝

def _read_csv(path: str) -> List[List[str]]:
    """Read CSV with auto-encoding detection and delimiter sniffing.

    Encoding order: utf-8-sig → utf-8 → cp1256 → windows-1252 → latin-1
    Delimiters tried: comma, semicolon, tab, pipe
    Falls back to utf-8 lossy if all strict encodings fail.
    """
    import csv, io as _io

    raw: Optional[str] = None
    used_enc = "utf-8"
    # windows-1252 is very common for Excel-exported Arabic CSVs on Windows
    for enc in ("utf-8-sig", "utf-8", "cp1256", "windows-1252", "latin-1"):
        try:
            with open(path, "r", encoding=enc, errors="strict") as fh:
                raw = fh.read()
            used_enc = enc
            break
        except (UnicodeDecodeError, UnicodeError):
            continue

    if raw is None:
        logger.warning("[csv] all strict encodings failed — using utf-8 lossy replacement")
        with open(path, "r", encoding="utf-8", errors="replace") as fh:
            raw = fh.read()
        used_enc = "utf-8(lossy)"

    # Manual BOM strip (shouldn't be needed with utf-8-sig but belt-and-suspenders)
    if raw.startswith('\ufeff'):
        raw = raw[1:]

    # ── Delimiter sniffing ────────────────────────────────────────────────────
    # Score each delimiter by: mean column count × (1 / coefficient_of_variation + 1)
    # High mean + low variation = consistent split = correct delimiter
    snippet = raw[:8192]
    best_delim, best_score = ",", 0.0
    delim_scores: Dict[str, str] = {}

    for d in (",", ";", "\t", "|"):
        try:
            reader = csv.reader(_io.StringIO(snippet), delimiter=d)
            sample: List[List[str]] = []
            for row in reader:
                if any(c.strip() for c in row):
                    sample.append(row)
                if len(sample) >= 12:
                    break
            if not sample:
                delim_scores[repr(d)] = "no_rows"
                continue
            widths = [len(r) for r in sample]
            if max(widths) <= 1:
                delim_scores[repr(d)] = "no_split"
                continue
            mean_w   = sum(widths) / len(widths)
            variance = sum((w - mean_w) ** 2 for w in widths) / len(widths)
            cv       = (variance ** 0.5) / (mean_w + 0.001)
            score    = mean_w / (cv + 1.0)
            delim_scores[repr(d)] = f"{score:.2f}"
            if score > best_score:
                best_score = score
                best_delim = d
        except Exception as ex:
            delim_scores[repr(d)] = f"err:{ex}"

    logger.info(f"[csv] enc={used_enc} delim_scores={delim_scores} → selected={best_delim!r}")

    rows = list(csv.reader(_io.StringIO(raw), delimiter=best_delim))
    clean_rows = [[_clean_cell(c) for c in r] for r in rows]
    logger.info(f"[csv] total_rows={len(clean_rows)}")
    return clean_rows


def _read_xlsx(path: str) -> List[List[List[str]]]:
    """Read all non-empty sheets from .xlsx using openpyxl.

    Returns list-of-sheets; each sheet is a list-of-rows; each row is a list of
    stripped string cells.  Leading and trailing all-empty rows are trimmed.
    """
    import openpyxl
    try:
        wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    except Exception as e:
        logger.error(f"[xlsx] openpyxl failed to open '{path}': {e}", exc_info=True)
        raise

    out: List[List[List[str]]] = []
    for ws in wb.worksheets:
        rows: List[List[str]] = []
        for row in ws.iter_rows():
            cells = [_clean_cell(str(c.value if c.value is not None else "")) for c in row]
            rows.append(cells)
        # Trim trailing all-empty rows
        while rows and not any(c for c in rows[-1]):
            rows.pop()
        # Trim leading all-empty rows
        while rows and not any(c for c in rows[0]):
            rows.pop(0)
        if rows:
            max_cols = max(len(r) for r in rows)
            out.append(rows)
            logger.info(f"[xlsx] sheet='{ws.title}' rows={len(rows)} max_cols={max_cols}")
        else:
            logger.debug(f"[xlsx] sheet='{ws.title}' is entirely empty — skipped")
    wb.close()

    if not out:
        logger.warning(f"[xlsx] workbook '{path}' has no non-empty sheets")
    return out


def _read_xls(path: str) -> List[List[List[str]]]:
    """Read all sheets from legacy .xls using xlrd.

    Falls back to openpyxl only if xlrd is not installed (openpyxl ≥ 3.x cannot
    read true .xls, so the fallback is a best-effort for xlsx-named-as-xls).
    """
    try:
        import xlrd
    except ImportError:
        logger.warning(
            "[xls] xlrd not installed — falling back to openpyxl (may fail for true .xls). "
            "Install: pip install xlrd>=2.0.0"
        )
        return _read_xlsx(path)

    try:
        wb = xlrd.open_workbook(path)
    except Exception as e:
        logger.error(f"[xls] xlrd failed to open '{path}': {e}", exc_info=True)
        raise

    out: List[List[List[str]]] = []
    for ws in wb.sheets():
        if ws.nrows == 0:
            logger.debug(f"[xls] sheet='{ws.name}' empty — skipped")
            continue
        rows: List[List[str]] = []
        for ri in range(ws.nrows):
            row: List[str] = []
            for ci in range(ws.ncols):
                ctype = ws.cell_type(ri, ci)
                val   = ws.cell_value(ri, ci)
                if ctype == xlrd.XL_CELL_NUMBER:
                    # xlrd returns all numbers as float; display whole numbers without ".0"
                    cell_str = (
                        str(int(val))
                        if (float(val) == int(val) and abs(val) < 1e15)
                        else str(val)
                    )
                elif ctype == xlrd.XL_CELL_DATE:
                    # Keep date cells as raw string; don't parse as price
                    cell_str = str(val)
                elif ctype == xlrd.XL_CELL_BOOLEAN:
                    cell_str = "نعم" if val else "لا"
                elif ctype == xlrd.XL_CELL_ERROR:
                    cell_str = ""
                else:
                    cell_str = str(val)
                row.append(_clean_cell(cell_str))
            rows.append(row)
        # Trim trailing / leading all-empty rows
        while rows and not any(c for c in rows[-1]):
            rows.pop()
        while rows and not any(c for c in rows[0]):
            rows.pop(0)
        if rows:
            out.append(rows)
            logger.info(f"[xls] sheet='{ws.name}' rows={len(rows)} cols={ws.ncols}")

    if not out:
        logger.warning(f"[xls] workbook '{path}' has no non-empty sheets")
    return out


# ╔══════════════════════════════════════════════════════════════════════════════╗
# ║                      SPREADSHEET ORCHESTRATOR                               ║
# ╚══════════════════════════════════════════════════════════════════════════════╝

def _score_sheet_quality(items: List[Dict]) -> float:
    """Score the quality of items extracted from one sheet (0.0–1.0).

    Combines four signals:
    - item count (saturates at 10+)
    - price coverage (fraction of items with a valid price)
    - average confidence
    - name quality (items with a meaningful, non-numeric name)

    Used to decide whether direct-parse results are production-worthy
    or we should fall back to OpenAI.
    """
    if not items:
        return 0.0
    n = len(items)
    if n < 2:
        return 0.15   # single item is suspicious — low but nonzero

    count_score  = min(n / 10.0, 1.0)
    price_score  = sum(1 for i in items if i.get("price") is not None) / n
    avg_conf     = sum(i.get("confidence", 0.0) for i in items) / n
    good_names   = sum(
        1 for i in items
        if len(i.get("name", "")) >= 2
        and not i.get("name", "").replace(" ", "").isdigit()
    )
    name_score = good_names / n

    score = (
        0.25 * count_score +
        0.30 * price_score +
        0.20 * avg_conf    +
        0.25 * name_score
    )
    return round(score, 3)


def _parse_spreadsheet(path: str, client, file_name: str = "") -> List[Dict]:
    """Parse .xlsx / .xls / .csv → product list.

    Flow:
    1. Route to correct reader based on extension.
    2. Read ALL sheets.
    3. Run direct column mapping on each sheet (_spreadsheet_direct).
    4. Score each sheet; keep sheets at ≥ 50% of the best sheet's score.
    5. Deduplicate items across sheets.
    6. If no items OR quality too low → fall back to OpenAI text parse.
       If OpenAI also returns nothing → return direct results anyway.

    Every major decision is logged for production diagnostics.
    """
    ext = Path(file_name).suffix.lower() if file_name else Path(path).suffix.lower()
    logger.info(f"[spreadsheet] ═══ START  file='{file_name or path}'  ext='{ext}' ═══")

    # ── Step 1: Read file ──────────────────────────────────────────────────────
    try:
        if ext == ".csv":
            logger.info("[spreadsheet] parser → csv reader")
            sheets: List[List[List[str]]] = [_read_csv(path)]
        elif ext == ".xlsx":
            logger.info("[spreadsheet] parser → openpyxl (.xlsx)")
            sheets = _read_xlsx(path)
        elif ext == ".xls":
            logger.info("[spreadsheet] parser → xlrd (.xls)")
            sheets = _read_xls(path)
        else:
            logger.warning(f"[spreadsheet] unknown ext='{ext}' — trying xlsx then csv")
            try:
                sheets = _read_xlsx(path)
                logger.info("[spreadsheet] sniff → openpyxl succeeded")
            except Exception:
                sheets = [_read_csv(path)]
                logger.info("[spreadsheet] sniff → csv reader succeeded")
    except Exception as e:
        logger.error(f"[spreadsheet] FATAL read error: {e}", exc_info=True)
        return _error_item(f"فشل قراءة الملف: {e}")

    sheets = [s for s in sheets if s]
    if not sheets:
        msg = "الملف فارغ أو لا يحتوي على بيانات قابلة للقراءة"
        logger.error(f"[spreadsheet] {msg}")
        return _error_item(msg)

    total_rows = sum(len(s) for s in sheets)
    logger.info(
        f"[spreadsheet] loaded  sheets={len(sheets)}  total_rows={total_rows}  "
        f"rows_per_sheet={[len(s) for s in sheets]}"
    )

    # ── Step 2: Direct mapping on each sheet ──────────────────────────────────
    sheet_results: List[Dict] = []
    for idx, rows in enumerate(sheets):
        label = f"sheet{idx + 1}"
        items = _spreadsheet_direct(rows, sheet_label=label)
        score = _score_sheet_quality(items)
        logger.info(
            f"[spreadsheet] {label}: direct_items={len(items)}  quality={score:.3f}"
        )
        sheet_results.append({"label": label, "items": items, "score": score})

    # ── Step 3: Collect sheets that meet quality bar ───────────────────────────
    best_score = max((r["score"] for r in sheet_results), default=0.0)
    all_items: List[Dict] = []
    used_sheets: List[str] = []

    for r in sheet_results:
        # Include this sheet only if:
        #   • it has items, AND
        #   • its score is at least 50% of the best sheet's score (avoids garbage sheets)
        quality_floor = max(best_score * 0.50, 0.10)
        if r["items"] and r["score"] >= quality_floor:
            all_items.extend(r["items"])
            used_sheets.append(r["label"])
        elif r["items"]:
            logger.info(
                f"[spreadsheet] {r['label']} excluded "
                f"(score={r['score']:.3f} < floor={quality_floor:.3f})"
            )

    logger.info(
        f"[spreadsheet] direct total: items={len(all_items)}  "
        f"from={used_sheets}  best_quality={best_score:.3f}"
    )

    # ── Step 4: Dedup across sheets (same item extracted from multiple sheets) ─
    dup_count = 0
    if len(used_sheets) > 1 and all_items:
        before = len(all_items)
        seen: Dict[str, Dict] = {}
        for item in all_items:
            key = _norm_header(item.get("name", ""))
            if key not in seen or item.get("confidence", 0) > seen[key].get("confidence", 0):
                seen[key] = item
        all_items = list(seen.values())
        dup_count = before - len(all_items)
        if dup_count:
            logger.info(f"[spreadsheet] cross-sheet dedup: removed {dup_count} duplicates")

    # ── Step 5: Decide whether to use direct results or fall back to OpenAI ───
    fallback_reason: Optional[str] = None
    if not all_items:
        fallback_reason = "no_items_extracted_from_any_sheet"
    elif best_score < _SHEET_QUALITY_FALLBACK:
        fallback_reason = (
            f"best_quality_score={best_score:.3f} < threshold={_SHEET_QUALITY_FALLBACK}"
        )

    if fallback_reason:
        logger.info(f"[spreadsheet] → OpenAI fallback  reason={fallback_reason}")
        text_parts: List[str] = []
        for idx, rows in enumerate(sheets):
            if len(sheets) > 1:
                text_parts.append(f"=== Sheet {idx + 1} ===")
            text_parts.extend(" | ".join(r) for r in rows[:300])
        text = "\n".join(text_parts)[:12000]
        openai_items = _call_text(text, client)
        logger.info(f"[spreadsheet] OpenAI returned {len(openai_items)} items")

        # If OpenAI also returns nothing, keep our low-quality direct results
        if not openai_items and all_items:
            logger.warning(
                "[spreadsheet] OpenAI returned empty — keeping low-quality direct results"
            )
            return all_items
        return openai_items

    logger.info(
        f"[spreadsheet] ═══ DONE  items={len(all_items)}  dups_removed={dup_count}  "
        f"openai_used=False ═══"
    )
    return all_items


# ╔══════════════════════════════════════════════════════════════════════════════╗
# ║                    DIRECT SPREADSHEET → PRODUCT MAPPER                      ║
# ╚══════════════════════════════════════════════════════════════════════════════╝

# Known food section / category keywords (normalised).
# Rows whose name matches one of these are treated as section-dividers and their
# name is carried forward as the running category for subsequent rows.
_SECTION_HEADER_PATTERNS = {
    # English
    "appetizer", "appetizers", "starter", "starters", "salad", "salads",
    "soup", "soups", "maincourse", "mainmenu", "main", "mains",
    "dessert", "desserts", "drink", "drinks", "beverage", "beverages",
    "sandwich", "sandwiches", "pizza", "burger", "burgers", "grill",
    "grills", "breakfast", "specials", "special", "kids", "kidsmenu",
    "sidedish", "sides", "combo", "combos", "set", "sets",
    # Arabic
    "مقبلات", "سلطات", "شوربات", "أطباقرئيسية", "وجباترئيسية",
    "حلويات", "مشروبات", "عصائر", "مشاوي", "بيتزا", "برغر",
    "سندويشات", "افطار", "وجباتخفيفة", "إضافات", "وجباتاطفال",
    "مقليات", "مخبوزات", "أرز", "معكرونة", "دجاج", "لحوم", "أسماك",
}


def _looks_like_section_header(name: str) -> bool:
    """Return True if the name matches a known food section keyword,
    or is a short ALL-CAPS ASCII string (English section divider style)."""
    hn = _norm_header(name)
    if hn in _SECTION_HEADER_PATTERNS:
        return True
    # Short ALL-CAPS ASCII name with no price → likely an English section header
    stripped = name.strip()
    if (
        len(stripped) <= 30
        and stripped == stripped.upper()
        and all(c.isascii() for c in stripped)
        and stripped.replace(' ', '').replace('-', '').replace('/', '').isalpha()
    ):
        return True
    return False


def _spreadsheet_direct(rows: List[List[str]], sheet_label: str = "") -> List[Dict]:
    """Map spreadsheet rows to product dicts using column headers.

    Algorithm:
    1. Scan the first 15 rows for the best header row (must include a name column;
       more recognised columns = better).
    2. If no header found → apply a no-header heuristic:
       a) Requires 2–5 consistent columns and ≥ 3 data rows.
       b) Identifies price column (most numeric values).
       c) Identifies name column (most text values, not price column).
       d) Requires ≥ 30% of rows to be numeric in the price column.
       e) Requires ≥ 50% of rows to be text in the name column.
    3. Delegate row conversion to _extract_data_rows.
    """
    prefix = f"[direct/{sheet_label}]" if sheet_label else "[direct]"
    if not rows:
        return []

    # ── Step 1: Find best header row (scan first 15 rows) ─────────────────────
    best_idx      = -1
    best_col_map: Dict[str, int] = {}
    best_mapped   = 0

    for i, row in enumerate(rows[:15]):
        if not any(c.strip() for c in row):
            continue
        col_map: Dict[str, int] = {}
        for ci, cell in enumerate(row):
            field = _header_to_field(_norm_header(cell))
            if field and field not in col_map:
                col_map[field] = ci
        if "name" in col_map and len(col_map) > best_mapped:
            best_mapped  = len(col_map)
            best_col_map = col_map
            best_idx     = i

    logger.info(f"{prefix} header_scan → idx={best_idx}  col_map={best_col_map}")

    if best_idx >= 0:
        data_rows = [r for r in rows[best_idx + 1:] if any(c.strip() for c in r)]
        if not data_rows:
            logger.warning(f"{prefix} header at row {best_idx} but zero data rows follow")
            return []
        return _extract_data_rows(rows, best_idx + 1, best_col_map, prefix, rows[best_idx], sheet_label)

    # ── Step 2: No-header heuristic ────────────────────────────────────────────
    logger.info(f"{prefix} no recognised header — attempting no-header heuristic")

    sample = [r for r in rows[:30] if any(c.strip() for c in r)]
    if len(sample) < 3:
        logger.warning(f"{prefix} only {len(sample)} non-empty rows — too few for heuristic")
        return []

    col_widths = [len([c for c in r if c.strip()]) for r in sample]
    min_w, max_w = min(col_widths), max(col_widths)

    if min_w < 2 or max_w > 6:
        logger.warning(f"{prefix} column width {min_w}–{max_w} — heuristic inconclusive")
        return []

    # Score each column for numeric (price) vs text (name) content
    price_hits = [0] * (max_w + 1)
    text_hits  = [0] * (max_w + 1)

    for r in sample:
        for ci, val in enumerate(r[:max_w + 1]):
            if not val.strip():
                continue
            if _parse_price(val) is not None:
                price_hits[ci] += 1
            else:
                text_hits[ci] += 1

    max_price_hits = max(price_hits) if price_hits else 0
    # Require ≥ 30% of sample rows to be numeric in some column
    if max_price_hits < max(len(sample) * 0.30, 2):
        logger.warning(
            f"{prefix} no column is sufficiently numeric "
            f"(max_price_hits={max_price_hits}/{len(sample)}) — heuristic skipped"
        )
        return []

    price_ci = max(range(len(price_hits)), key=lambda i: price_hits[i])

    # Name column: best text-rich column that is NOT the price column
    text_hits_excl = [
        (text_hits[i] if i != price_ci else -1)
        for i in range(len(text_hits))
    ]
    name_ci = max(range(len(text_hits_excl)), key=lambda i: text_hits_excl[i])

    # Require ≥ 50% text values in the name column
    name_text_frac = text_hits[name_ci] / max(len(sample), 1)
    if name_text_frac < 0.50:
        logger.warning(
            f"{prefix} name column has only {name_text_frac:.0%} text — heuristic skipped"
        )
        return []

    cat_ci: Optional[int] = None
    if max_w >= 3:
        rem = [i for i in range(min(max_w, len(text_hits))) if i not in (name_ci, price_ci)]
        if rem:
            cat_ci = max(rem, key=lambda i: text_hits[i])

    heuristic_map: Dict[str, int] = {"name": name_ci, "price": price_ci}
    if cat_ci is not None:
        heuristic_map["category"] = cat_ci

    logger.info(f"{prefix} heuristic col_map={heuristic_map}")
    return _extract_data_rows(rows, 0, heuristic_map, prefix, header_row=None, sheet_label=sheet_label)


# ── Row extraction (core data loop) ───────────────────────────────────────────

def _extract_data_rows(
    rows: List[List[str]],
    start_idx: int,
    col_map: Dict[str, int],
    prefix: str,
    header_row: Optional[List[str]],
    sheet_label: str = "",
) -> List[Dict]:
    """Convert spreadsheet data rows to product dicts.

    Skips:
    - Blank rows
    - Decoration rows (---, ===, ***)
    - Pure-numeric names (likely row indices)
    - Repeated header rows
    - Known total / footer / metadata rows (_SKIP_ROW_NAMES)
    - Section-header rows (carried forward as implicit category instead)

    Scores each item with a 4-factor confidence:
    - price present (weight 0.50)
    - name length quality (weight 0.25)
    - category present (weight 0.15)
    - description present (weight 0.10)

    Adds source_sheet and source_row to every item for preview traceability.
    """
    items:   List[Dict] = []
    skipped: int = 0
    skip_reasons: Dict[str, int] = {}

    def _bump(reason: str) -> None:
        nonlocal skipped
        skipped += 1
        skip_reasons[reason] = skip_reasons.get(reason, 0) + 1

    # Normalised header tokens (for repeated-header row detection)
    header_norm: set = set()
    if header_row:
        for c in header_row:
            hn = _norm_header(c)
            if hn:
                header_norm.add(hn)

    # Running category context (updated by section-header rows)
    current_category = "عام"

    for row_offset, row in enumerate(rows[start_idx:]):
        abs_row = start_idx + row_offset + 2  # 1-based, +1 for header, +1 for 0-index

        # ── Blank row ──────────────────────────────────────────────────────────
        if not row or not any(c.strip() for c in row):
            _bump("blank")
            continue

        # ── Cell accessor with closure capture ────────────────────────────────
        def _cell(key: str, default: str = "", _cm=col_map, _r=row) -> str:
            i = _cm.get(key)
            if i is None or i >= len(_r):
                return default
            return _clean_cell(str(_r[i]))

        name = _cell("name")

        # ── Short / empty name ─────────────────────────────────────────────────
        if not name or len(name) < 2:
            _bump("name_too_short")
            continue

        # ── Decoration rows (lines of dashes, equals, asterisks, underscores) ─
        stripped_decoration = (
            name.replace(' ', '').replace('-', '').replace('=', '')
                .replace('*', '').replace('#', '').replace('_', '')
                .replace('~', '').replace('.', '')
        )
        if not stripped_decoration:
            _bump("decoration")
            continue

        # ── Pure-numeric name (row index like 1, 2, 3, …) ─────────────────────
        if name.replace(' ', '').replace('.', '').replace(',', '').isdigit():
            _bump("numeric_name")
            continue

        name_norm = _norm_header(name)

        # ── Repeated header row ────────────────────────────────────────────────
        if header_norm and name_norm in header_norm:
            _bump("repeated_header")
            continue

        # ── Known total / footer / metadata patterns ───────────────────────────
        if name_norm in _SKIP_ROW_NAMES:
            _bump("skip_pattern")
            continue

        # ── Section-header carry-forward ───────────────────────────────────────
        # If this row has no price AND looks like a section header,
        # update running category and skip it as a product row.
        price_raw   = _cell("price")
        parsed_price = _parse_price(price_raw) if price_raw else None

        if parsed_price is None and _looks_like_section_header(name):
            current_category = name
            logger.debug(f"{prefix} row={abs_row} section_header → category='{name}'")
            _bump("section_header")
            continue

        # ── Build product item ─────────────────────────────────────────────────
        cat_raw  = _cell("category")
        category = cat_raw if cat_raw else current_category

        desc = _cell("description")

        # Multi-factor confidence score
        conf = round(
            0.50 * (1.0 if parsed_price is not None else 0.0)
            + 0.25 * min(len(name) / 8.0, 1.0)          # ≥ 8 chars → full weight
            + 0.15 * (1.0 if category and category != "عام" else 0.5)
            + 0.10 * (1.0 if desc else 0.7),
            2,
        )
        needs_review = conf < REVIEW_THRESHOLD or parsed_price is None

        # Build human-readable source note
        notes: List[str] = []
        if price_raw and parsed_price is None:
            notes.append(f"سعر غير صالح: '{price_raw}'")
        elif parsed_price is None:
            notes.append("السعر مفقود")

        items.append({
            "name":         name,
            "category":     category,
            "price":        parsed_price,
            "description":  desc,
            "variants":     [],
            "confidence":   conf,
            "needs_review": needs_review,
            "source_note":  "; ".join(notes),
            "source_sheet": sheet_label,
            "source_row":   abs_row,
        })

    logger.info(
        f"{prefix} extracted={len(items)}  skipped={skipped}  "
        f"skip_reasons={skip_reasons}"
    )
    return items


def _error_item(msg: str) -> List[Dict]:
    return [{
        "name": "خطأ في القراءة", "category": "", "price": None,
        "description": msg, "variants": [], "confidence": 0.0,
        "needs_review": True, "source_note": msg
    }]


# ── Normalization ─────────────────────────────────────────────────────────────

def _normalize(raw_items: List[Dict], session_id: str) -> List[Dict]:
    """
    Normalize raw OpenAI output:
    - Add temp_id, session_id, action
    - Coerce types, cap confidence
    - Flag items needing review
    - Deduplicate by name (keep highest confidence)
    - Preserve source_sheet / source_row from direct-parse items
    """
    by_name: Dict[str, Dict] = {}  # name_lower → item

    for raw in raw_items:
        if not isinstance(raw, dict):
            continue

        name = str(raw.get("name") or "").strip()
        if not name or len(name) < 2:
            continue

        # Price coercion
        price = raw.get("price")
        if price is not None:
            try:
                price = float(str(price).replace(",", ".").replace(" ", ""))
                price = price if price >= 0 else None
            except (ValueError, TypeError):
                price = None

        # Confidence coercion
        try:
            conf = float(raw.get("confidence", 0.8))
            conf = max(0.0, min(1.0, conf))
        except (ValueError, TypeError):
            conf = 0.5

        # Lower confidence when price is missing
        if price is None:
            conf = min(conf, 0.55)

        needs_review = bool(raw.get("needs_review")) or conf < REVIEW_THRESHOLD or price is None

        # Normalize variants
        variants = raw.get("variants") or []
        clean_v = []
        for v in (variants if isinstance(variants, list) else []):
            if isinstance(v, dict) and str(v.get("name", "")).strip():
                vp = v.get("price")
                try:
                    vp = float(vp) if vp is not None else None
                except (ValueError, TypeError):
                    vp = None
                clean_v.append({"name": str(v["name"]).strip(), "price": vp})

        item = {
            "temp_id":      str(uuid.uuid4()),
            "session_id":   session_id,
            "name":         name,
            "category":     str(raw.get("category") or "عام").strip() or "عام",
            "price":        price,
            "description":  str(raw.get("description") or "").strip(),
            "variants":     clean_v,
            "confidence":   round(conf, 2),
            "needs_review": needs_review,
            "source_note":  str(raw.get("source_note") or "").strip(),
            "source_page":  int(raw.get("source_page") or 1),
            "source_sheet": str(raw.get("source_sheet") or ""),
            "source_row":   int(raw.get("source_row") or 0),
            "action":       "create",   # create | skip | update_existing
            "existing_match": None,
            "image_url":    "",
        }

        name_key = name.lower()
        if name_key not in by_name or conf > by_name[name_key]["confidence"]:
            by_name[name_key] = item

    result = list(by_name.values())
    logger.info(f"_normalize: input={len(raw_items)} raw → output={len(result)} unique items")
    return result


# ── Duplicate detection ────────────────────────────────────────────────────────

def _name_similarity(a: str, b: str) -> float:
    """Simple overlap similarity between two strings."""
    a, b = a.lower().strip(), b.lower().strip()
    if a == b:
        return 1.0
    if len(a) > 3 and (a in b or b in a):
        return 0.90
    # Bigram overlap
    def bigrams(s):
        return set(s[i:i+2] for i in range(len(s) - 1))
    ab, bb = bigrams(a), bigrams(b)
    if not ab or not bb:
        return 0.0
    return 2 * len(ab & bb) / (len(ab) + len(bb))


def detect_duplicates(items: List[Dict], restaurant_id: str) -> List[Dict]:
    """
    For each extracted item, check against existing products in the DB.
    Sets item["existing_match"] and item["action"] = "merge_review" when
    a likely duplicate (similarity ≥ 0.80) is found.
    """
    import database
    conn = database.get_db()
    try:
        existing = conn.execute(
            "SELECT id, name, category, price FROM products WHERE restaurant_id=?",
            (restaurant_id,)
        ).fetchall()
    finally:
        conn.close()

    if not existing:
        return items

    ex_list = [dict(e) for e in existing]

    for item in items:
        best_score  = 0.0
        best_match  = None
        for ex in ex_list:
            score = _name_similarity(item["name"], ex["name"])
            if score > best_score:
                best_score = score
                best_match = ex

        if best_score >= 0.80 and best_match:
            item["existing_match"] = {
                "id":         best_match["id"],
                "name":       best_match["name"],
                "price":      best_match["price"],
                "similarity": round(best_score, 2),
            }
            item["action"] = "merge_review"  # user must decide

    return items


# ── Main entry points ─────────────────────────────────────────────────────────

def parse_files(file_paths: List[str], file_names: List[str], session_id: str) -> List[Dict]:
    """
    Parse one or more menu files. Images are batched together (up to 4 per call)
    for efficiency. Non-image files are processed individually.

    Returns a normalized, deduplicated list of menu items.
    """
    client = _get_client()

    image_batches: List[List[str]] = []
    cur_batch: List[str] = []
    non_image_items: List[Dict] = []

    for path, name in zip(file_paths, file_names):
        ext = Path(name).suffix.lower()
        if ext in SUPPORTED_IMAGES:
            cur_batch.append(path)
            if len(cur_batch) >= 4:
                image_batches.append(cur_batch)
                cur_batch = []
        else:
            # Process non-image file — quota errors propagate up
            raw = _dispatch_non_image(path, name, client)
            non_image_items.extend(raw)

    if cur_batch:
        image_batches.append(cur_batch)

    # Call Vision for each batch of images
    image_raw: List[Dict] = []
    for batch in image_batches:
        image_raw.extend(_call_vision(batch, client))

    all_raw = image_raw + non_image_items
    logger.info(f"session={session_id}: {len(all_raw)} raw items from {len(file_paths)} files")

    return _normalize(all_raw, session_id)


def _dispatch_non_image(path: str, name: str, client) -> List[Dict]:
    ext = Path(name).suffix.lower()
    logger.info(f"[dispatch] file='{name}' ext='{ext}' path='{path}'")
    if ext == ".pdf":
        return _parse_pdf(path, client)
    elif ext == ".docx":
        return _parse_docx(path, client)
    elif ext in {".xlsx", ".xls", ".csv"}:
        logger.info(f"[dispatch] → _parse_spreadsheet (direct column mapping, no OpenAI)")
        return _parse_spreadsheet(path, client, file_name=name)
    elif ext == ".txt":
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            return _call_text(f.read(), client)
    else:
        raise ValueError(f"Unsupported file type: {ext}")
